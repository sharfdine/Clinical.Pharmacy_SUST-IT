"""
report_generator.py — Word document generator for clinical pharmacy monthly reports.

Creates a professional .docx report with:
  1. Title page
  2. Introduction
  3. Findings (with embedded charts and tables)
  4. Discussion
  5. Conclusion
"""

import os
import tempfile
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import plotly.graph_objects as go
import pandas as pd


# ---------------------------------------------------------------------------
# Styles / constants
# ---------------------------------------------------------------------------
HEADING_COLOR = RGBColor(0x1A, 0x1A, 0x2E)  # Dark navy
ACCENT_COLOR = RGBColor(0x4F, 0x8F, 0xEA)   # Blue
TEXT_COLOR = RGBColor(0x33, 0x33, 0x33)       # Dark gray


def _set_cell_shading(cell, color_hex: str):
    """Set background color for a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._element.get_or_add_tcPr().append(shading)


def _add_styled_table(doc, df: pd.DataFrame, header_color: str = "4F8FEA"):
    """Add a formatted table to the document from a DataFrame."""
    if df.empty:
        doc.add_paragraph("No data available for this section.", style="Body Text")
        return

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    # Header row
    header_row = table.rows[0]
    for i, col_name in enumerate(df.columns):
        cell = header_row.cells[i]
        cell.text = str(col_name)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(cell, header_color)

    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph("")  # spacing


def _save_chart_as_image(fig: go.Figure, temp_dir: str, name: str) -> str:
    """Export a Plotly figure to a PNG image file and return the path."""
    path = os.path.join(temp_dir, f"{name}.png")
    fig.write_image(path, width=900, height=450, scale=2)
    return path


def _add_chart_to_doc(doc, fig: go.Figure, temp_dir: str, name: str, width: float = 6.0):
    """Save chart as image and insert into the Word document."""
    img_path = _save_chart_as_image(fig, temp_dir, name)
    doc.add_picture(img_path, width=Inches(width))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")  # spacing


def generate_report(
    month_year: str,
    hospital_name: str,
    intervention_results: dict | None,
    ward_df: pd.DataFrame | None,
    consultant_df: pd.DataFrame | None,
    error_results: dict | None,
    error_ward_df: pd.DataFrame | None,
    ham_lasa_results: dict | None,
    ham_patient_df: pd.DataFrame | None,
    adr_results: dict | None,
    charts: dict[str, go.Figure] | None,
) -> BytesIO:
    """
    Generate a complete Word report and return it as a BytesIO buffer.

    Parameters
    ----------
    month_year : str
        Report month, e.g. "July 2026".
    hospital_name : str
        Name of the hospital for the title page.
    intervention_results, ward_df, consultant_df, error_results,
    error_ward_df, ham_lasa_results, adr_results : analysis outputs
    charts : dict[str, go.Figure]
        Named chart figures to embed.

    Returns
    -------
    BytesIO
        Word document in memory.
    """
    doc = Document()
    temp_dir = tempfile.mkdtemp()

    # ── Page Setup ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title Page ─────────────────────────────────────────────────────
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(hospital_name.upper())
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = HEADING_COLOR

    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Clinical Pharmacy Department")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = ACCENT_COLOR

    doc.add_paragraph("")

    report_title = doc.add_paragraph()
    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = report_title.add_run(f"Monthly Report — {month_year}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph("")
    doc.add_paragraph("")

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    doc.add_page_break()

    # ── Table of Contents placeholder ──────────────────────────────────
    toc_heading = doc.add_heading("Table of Contents", level=1)
    for run in toc_heading.runs:
        run.font.color.rgb = HEADING_COLOR

    toc_items = [
        "1. Introduction",
        "2. Findings",
        "   2.1 Pharmacist Interventions",
        "   2.2 Medication Errors",
        "   2.3 HAM/LASA Consumption",
        "   2.4 Adverse Drug Reactions",
        "3. Discussion",
        "4. Conclusion",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ── 1. Introduction ───────────────────────────────────────────────
    heading = doc.add_heading("1. Introduction", level=1)
    for run in heading.runs:
        run.font.color.rgb = HEADING_COLOR

    intro_text = (
        f"This report presents the monthly analysis of clinical pharmacy activities at "
        f"{hospital_name} for the month of {month_year}. The Clinical Pharmacy Department "
        f"plays a vital role in ensuring medication safety, optimizing therapeutic outcomes, "
        f"and reducing medication-related errors across all hospital departments.\n\n"
        f"The report covers the following key areas:\n"
    )
    p = doc.add_paragraph(intro_text, style="Body Text")
    for run in p.runs:
        run.font.size = Pt(11)

    intro_bullets = [
        "Pharmacist interventions during prescription review and their acceptance rates",
        "Medication errors identified and classified by type and ward",
        "Consumption patterns of High Alert Medications (HAM) and Look-Alike Sound-Alike (LASA) drugs",
        "Adverse Drug Reactions (ADRs) reported, including those associated with HAM/LASA medications",
    ]
    for bullet in intro_bullets:
        bp = doc.add_paragraph(bullet, style="List Bullet")
        for run in bp.runs:
            run.font.size = Pt(11)

    doc.add_paragraph("")

    # ── 2. Findings ───────────────────────────────────────────────────
    heading = doc.add_heading("2. Findings", level=1)
    for run in heading.runs:
        run.font.color.rgb = HEADING_COLOR

    # ── 2.1 Interventions ─────────────────────────────────────────────
    heading = doc.add_heading("2.1 Pharmacist Interventions", level=2)
    for run in heading.runs:
        run.font.color.rgb = ACCENT_COLOR

    if intervention_results:
        total_files = intervention_results["total_files_reviewed"]
        total_int = intervention_results["total_interventions"]
        accepted = intervention_results["accepted"]
        rejected = intervention_results["rejected"]
        rate = intervention_results["acceptance_rate"]

        p = doc.add_paragraph(
            f"During {month_year}, a total of {total_files} patient files were reviewed by the "
            f"clinical pharmacists. A total of {total_int} interventions were made, of which "
            f"{accepted} ({rate}%) were accepted and {rejected} were rejected by the prescribing physicians.",
            style="Body Text",
        )
        for run in p.runs:
            run.font.size = Pt(11)

        # Summary table
        summary_df = pd.DataFrame({
            "Metric": ["Files Reviewed", "Total Interventions", "Accepted", "Rejected", "Acceptance Rate"],
            "Value": [total_files, total_int, accepted, rejected, f"{rate}%"],
        })
        _add_styled_table(doc, summary_df)

        # Acceptance pie chart
        if charts and "acceptance_pie" in charts:
            _add_chart_to_doc(doc, charts["acceptance_pie"], temp_dir, "acceptance_pie")
            cap = doc.add_paragraph("Figure 1: Intervention Acceptance Rate")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)
                run.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            doc.add_paragraph("")

    # Ward-wise interventions
    if ward_df is not None and not ward_df.empty:
        p = doc.add_paragraph(
            f"The ward-wise distribution of interventions shows that {ward_df.iloc[0]['Ward']} had the "
            f"highest number of interventions ({ward_df.iloc[0]['Count']}), followed by "
            f"{ward_df.iloc[1]['Ward'] if len(ward_df) > 1 else 'N/A'} "
            f"({ward_df.iloc[1]['Count'] if len(ward_df) > 1 else 'N/A'}).",
            style="Body Text",
        )
        for run in p.runs:
            run.font.size = Pt(11)

        _add_styled_table(doc, ward_df.head(10))

        if charts and "ward_interventions" in charts:
            _add_chart_to_doc(doc, charts["ward_interventions"], temp_dir, "ward_interventions")
            cap = doc.add_paragraph("Figure 2: Interventions by Ward")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)
                run.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            doc.add_paragraph("")

    # Consultant-wise interventions
    if consultant_df is not None and not consultant_df.empty:
        p = doc.add_paragraph(
            f"The consultant-wise analysis reveals that {consultant_df.iloc[0]['Consultant']} received the "
            f"most interventions ({consultant_df.iloc[0]['Count']}). The top consultants receiving "
            f"interventions are listed below.",
            style="Body Text",
        )
        for run in p.runs:
            run.font.size = Pt(11)

        _add_styled_table(doc, consultant_df.head(10))

        if charts and "consultant" in charts:
            _add_chart_to_doc(doc, charts["consultant"], temp_dir, "consultant")
            cap = doc.add_paragraph("Figure 3: Interventions by Consultant")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)
                run.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            doc.add_paragraph("")

    # ── 2.2 Medication Errors ─────────────────────────────────────────
    heading = doc.add_heading("2.2 Medication Errors", level=2)
    for run in heading.runs:
        run.font.color.rgb = ACCENT_COLOR

    if error_results and error_results["total_errors"] > 0:
        total_err = error_results["total_errors"]
        p = doc.add_paragraph(
            f"A total of {total_err} medication errors were identified and recorded during {month_year}. "
            f"The errors were classified into the following categories:",
            style="Body Text",
        )
        for run in p.runs:
            run.font.size = Pt(11)

        # Error breakdown table
        err_summary = pd.DataFrame({
            "Error Type": ["Prescription Errors", "Administration Errors", "Transcription Errors",
                          "Illegible Handwriting", "Incorrect Abbreviation", "Other"],
            "Count": [error_results["prescription"], error_results["administration"],
                     error_results["transcription"], error_results["illegible_handwriting"],
                     error_results["incorrect_abbreviation"], error_results["other"]],
        })
        err_summary = err_summary[err_summary["Count"] > 0]
        _add_styled_table(doc, err_summary)

        if charts and "error_type" in charts:
            _add_chart_to_doc(doc, charts["error_type"], temp_dir, "error_type")
            cap = doc.add_paragraph("Figure 4: Medication Error Types Distribution")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)
                run.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            doc.add_paragraph("")

        # Errors by ward
        if error_ward_df is not None and not error_ward_df.empty:
            p = doc.add_paragraph(
                "The following table and chart show the distribution of errors across hospital wards:",
                style="Body Text",
            )
            for run in p.runs:
                run.font.size = Pt(11)

            ward_err_table = error_ward_df.reset_index()
            _add_styled_table(doc, ward_err_table)

            if charts and "errors_by_ward" in charts:
                _add_chart_to_doc(doc, charts["errors_by_ward"], temp_dir, "errors_by_ward")
                cap = doc.add_paragraph("Figure 5: Errors by Ward")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.font.size = Pt(9)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                doc.add_paragraph("")
    else:
        p = doc.add_paragraph(
            f"No medication errors were recorded during {month_year}.",
            style="Body Text",
        )
        for run in p.runs:
            run.font.size = Pt(11)

    # ── 2.3 HAM/LASA ─────────────────────────────────────────────────
    heading = doc.add_heading("2.3 HAM/LASA Consumption", level=2)
    for run in heading.runs:
        run.font.color.rgb = ACCENT_COLOR

    if ham_lasa_results and ham_lasa_results["total_records"] > 0:
        total_pts = ham_lasa_results["total_patients"]
        total_recs = ham_lasa_results["total_records"]

        p = doc.add_paragraph(
            f"During {month_year}, a total of {total_pts} patients received High Alert Medications (HAM) "
            f"or Look-Alike Sound-Alike (LASA) drugs, with {total_recs} total consumption records.",
            style="Body Text",
        )
        for run in p.runs:
            run.font.size = Pt(11)

        # HAM vs LASA split
        if not ham_lasa_results["by_type"].empty:
            _add_styled_table(doc, ham_lasa_results["by_type"])

            if charts and "ham_lasa_type" in charts:
                _add_chart_to_doc(doc, charts["ham_lasa_type"], temp_dir, "ham_lasa_type")
                cap = doc.add_paragraph("Figure 6: HAM vs LASA Distribution")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.font.size = Pt(9)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                doc.add_paragraph("")

        # Most common drugs
        if not ham_lasa_results["by_drug"].empty:
            top_drug = ham_lasa_results["by_drug"].iloc[0]
            p = doc.add_paragraph(
                f"The most commonly administered HAM/LASA medication was {top_drug['Drug']} "
                f"with a consumption count of {top_drug['Count']}.",
                style="Body Text",
            )
            for run in p.runs:
                run.font.size = Pt(11)

            _add_styled_table(doc, ham_lasa_results["by_drug"].head(10))

            if charts and "ham_lasa_drugs" in charts:
                _add_chart_to_doc(doc, charts["ham_lasa_drugs"], temp_dir, "ham_lasa_drugs")
                cap = doc.add_paragraph("Figure 7: Most Common HAM/LASA Medications")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.font.size = Pt(9)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                doc.add_paragraph("")

        # Frequency by patient
        if ham_patient_df is not None and not ham_patient_df.empty:
            top_patient = ham_patient_df.iloc[0]
            p = doc.add_paragraph(
                f"{top_patient['Patient Name']} (MR No {top_patient['MR No']}) received the most "
                f"HAM/LASA doses during the month, with {top_patient['Count']} recorded instances. "
                f"The frequency of HAM/LASA administration per patient is summarized below.",
                style="Body Text",
            )
            for run in p.runs:
                run.font.size = Pt(11)

            _add_styled_table(doc, ham_patient_df.head(10))

            if charts and "ham_lasa_patient" in charts:
                _add_chart_to_doc(doc, charts["ham_lasa_patient"], temp_dir, "ham_lasa_patient")
                cap = doc.add_paragraph("Figure 8: HAM/LASA Frequency by Patient")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.font.size = Pt(9)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                doc.add_paragraph("")
    else:
        p = doc.add_paragraph(
            f"No HAM/LASA consumption records were found for {month_year}.",
            style="Body Text",
        )
        for run in p.runs:
            run.font.size = Pt(11)

    # ── 2.4 ADR ───────────────────────────────────────────────────────
    heading = doc.add_heading("2.4 Adverse Drug Reactions (ADRs)", level=2)
    for run in heading.runs:
        run.font.color.rgb = ACCENT_COLOR

    if adr_results and adr_results["total_adrs"] > 0:
        total_adrs = adr_results["total_adrs"]
        from_ham = adr_results["adrs_from_ham_lasa"]
        other = adr_results["adrs_other"]

        p = doc.add_paragraph(
            f"A total of {total_adrs} Adverse Drug Reactions were reported during {month_year}. "
            f"Of these, {from_ham} ({round(from_ham/total_adrs*100, 1) if total_adrs > 0 else 0}%) "
            f"were associated with HAM/LASA medications, while {other} were from other drugs.",
            style="Body Text",
        )
        for run in p.runs:
            run.font.size = Pt(11)

        adr_summary = pd.DataFrame({
            "Category": ["Total ADRs", "From HAM/LASA", "From Other Drugs"],
            "Count": [total_adrs, from_ham, other],
        })
        _add_styled_table(doc, adr_summary)

        if charts and "adr_summary" in charts:
            _add_chart_to_doc(doc, charts["adr_summary"], temp_dir, "adr_summary")
            cap = doc.add_paragraph("Figure 9: ADR Source Distribution")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)
                run.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            doc.add_paragraph("")

        if not adr_results["by_drug"].empty:
            p = doc.add_paragraph(
                "The following table shows the drugs associated with the reported ADRs:",
                style="Body Text",
            )
            for run in p.runs:
                run.font.size = Pt(11)
            _add_styled_table(doc, adr_results["by_drug"].head(10))
    else:
        p = doc.add_paragraph(
            f"No Adverse Drug Reactions were reported during {month_year}.",
            style="Body Text",
        )
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ── 3. Discussion ─────────────────────────────────────────────────
    heading = doc.add_heading("3. Discussion", level=1)
    for run in heading.runs:
        run.font.color.rgb = HEADING_COLOR

    discussion_points = []

    if intervention_results:
        rate = intervention_results["acceptance_rate"]
        if rate >= 80:
            discussion_points.append(
                f"The intervention acceptance rate of {rate}% indicates a strong collaborative "
                f"relationship between the clinical pharmacists and the medical team. This reflects "
                f"the high quality and clinical relevance of the interventions made."
            )
        elif rate >= 50:
            discussion_points.append(
                f"The intervention acceptance rate of {rate}% suggests moderate acceptance of "
                f"pharmacist recommendations. Strategies to improve communication and evidence-based "
                f"recommendations should be explored to enhance acceptance rates."
            )
        else:
            discussion_points.append(
                f"The intervention acceptance rate of {rate}% indicates room for improvement. "
                f"A review of intervention strategies and enhanced collaboration with prescribers "
                f"is recommended."
            )

    if error_results and error_results["total_errors"] > 0:
        max_type = max(
            ["prescription", "administration", "transcription", "illegible_handwriting", "incorrect_abbreviation"],
            key=lambda x: error_results.get(x, 0),
        )
        max_count = error_results[max_type]
        discussion_points.append(
            f"Among medication errors, {max_type.replace('_', ' ').title()} errors were the most "
            f"prevalent ({max_count} cases). This highlights the need for targeted interventions "
            f"and training programs focused on reducing these specific error types."
        )

    if ham_lasa_results and ham_lasa_results["total_records"] > 0:
        discussion_points.append(
            f"The consumption of HAM/LASA medications requires continued vigilance. Proper "
            f"labeling, storage, and dispensing protocols should be regularly audited to "
            f"minimize the risk of errors associated with these high-risk medications."
        )

    if adr_results and adr_results["total_adrs"] > 0:
        from_ham = adr_results["adrs_from_ham_lasa"]
        total = adr_results["total_adrs"]
        if from_ham > 0:
            discussion_points.append(
                f"The occurrence of {from_ham} ADR(s) from HAM/LASA medications "
                f"({round(from_ham/total*100, 1)}% of all ADRs) underscores the importance of "
                f"pharmacovigilance activities and close monitoring of patients receiving "
                f"high-alert medications."
            )

    if not discussion_points:
        discussion_points.append(
            "The clinical pharmacy activities during this period demonstrate the department's "
            "commitment to medication safety and patient care optimization."
        )

    for point in discussion_points:
        p = doc.add_paragraph(point, style="Body Text")
        for run in p.runs:
            run.font.size = Pt(11)
        doc.add_paragraph("")

    # ── 4. Conclusion ─────────────────────────────────────────────────
    heading = doc.add_heading("4. Conclusion", level=1)
    for run in heading.runs:
        run.font.color.rgb = HEADING_COLOR

    conclusion_parts = []

    conclusion_parts.append(
        f"The clinical pharmacy department's activities during {month_year} demonstrate "
        f"a continued commitment to enhancing medication safety and patient outcomes at "
        f"{hospital_name}."
    )

    if intervention_results:
        conclusion_parts.append(
            f"A total of {intervention_results['total_interventions']} interventions were made "
            f"with an acceptance rate of {intervention_results['acceptance_rate']}%."
        )

    if error_results and error_results["total_errors"] > 0:
        conclusion_parts.append(
            f"{error_results['total_errors']} medication errors were identified and classified, "
            f"contributing to the hospital's quality improvement initiatives."
        )

    conclusion_parts.append(
        "The department recommends the following actions for the upcoming month:"
    )

    for part in conclusion_parts:
        p = doc.add_paragraph(part, style="Body Text")
        for run in p.runs:
            run.font.size = Pt(11)

    recommendations = [
        "Continue regular prescription review and documentation of interventions",
        "Conduct targeted training sessions on the most common error types identified",
        "Strengthen the HAM/LASA medication safety protocols",
        "Enhance ADR reporting and pharmacovigilance activities",
        "Promote interdisciplinary collaboration for better patient outcomes",
    ]
    for rec in recommendations:
        bp = doc.add_paragraph(rec, style="List Bullet")
        for run in bp.runs:
            run.font.size = Pt(11)

    # ── Save to BytesIO ───────────────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Clean up temp images
    for f in os.listdir(temp_dir):
        try:
            os.remove(os.path.join(temp_dir, f))
        except OSError:
            pass
    try:
        os.rmdir(temp_dir)
    except OSError:
        pass

    return buffer
